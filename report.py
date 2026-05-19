import ollama
from docx import Document

topic = input("Enter Topic: ")

response = ollama.chat(
    model='phi3:mini',
    messages=[
        {
            'role': 'user',
            'content': f'Write a detailed report on {topic}'
        }
    ]
)

report = response['message']['content']

doc = Document()
doc.add_heading(topic, 0)
doc.add_paragraph(report)

filename = f"{topic}.docx"
doc.save(filename)

print("Word file created successfully!")
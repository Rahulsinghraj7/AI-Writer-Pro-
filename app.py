import streamlit as st
import ollama
from docx import Document

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="AI Writer Pro",
    layout="centered"
)

# ---------------- MAIN TITLE ----------------
st.title("⚡ AI Writer Pro")
st.caption("Live AI Document Generator with DOCX + PDF Export")

# ---------------- SIDEBAR ----------------
with st.sidebar:

    st.title("⚙️ AI Writer Pro")

    st.subheader("Document Settings")

    # Topic
    topic = st.text_input("📌 Enter Topic")

    # Document Type
    mode = st.selectbox(
        "📝 Select Document Type",
        ["Report", "Resume", "Blog", "Notes", "Story"]
    )

    # Tone
    tone = st.selectbox(
        "🎯 Select Tone",
        ["Professional", "Simple", "Formal", "Creative"]
    )

    # Pages
    pages = st.slider(
        "📄 Number of Pages",
        min_value=1,
        max_value=5,
        value=2
    )

    # Word Calculation
    word_limit = pages * 120

    st.info(f"Estimated Output: ~{word_limit} words")

    st.markdown("---")

    st.success("AI Writer Pro Live Generator")

    # Generate Button
    generate_btn = st.button("🚀 Generate Document")

# ---------------- TASK MAP ----------------
task_map = {
    "Report": "Write a professional report",
    "Resume": "Create an ATS-friendly resume",
    "Blog": "Write an SEO-friendly blog article",
    "Notes": "Write clean exam notes",
    "Story": "Write an engaging creative story with characters, emotions, dialogues, and a meaningful ending"
}

task = task_map[mode]

# ---------------- EXTRA INSTRUCTIONS ----------------
extra_instruction = ""

# STORY
if mode == "Story":
    extra_instruction = """
    Include:
    - Character names
    - Interesting plot
    - Dialogues
    - Emotional storytelling
    - Proper ending
    """

# RESUME
# ---------------- RESUME INPUTS ----------------

resume_data = ""

if mode == "Resume":

    st.subheader("📄 Resume Details")

    name = st.text_input("Full Name")

    email = st.text_input("Email")

    phone = st.text_input("Phone Number")

    skills = st.text_area("Skills")

    education = st.text_area("Education")

    projects = st.text_area("Projects")

    experience = st.text_area("Experience")

    certifications = st.text_area("Certifications")

    # Store Everything
    resume_data = f"""
    Name: {name}
    Email: {email}
    Phone: {phone}

    Skills:
    {skills}

    Education:
    {education}

    Projects:
    {projects}

    Experience:
    {experience}

    Certifications:
    {certifications}
    """

# BLOG
elif mode == "Blog":
    extra_instruction = """
    Include:
    - Catchy title
    - SEO-friendly headings
    - Introduction
    - Main content
    - Bullet points
    - Conclusion
    """

# NOTES
elif mode == "Notes":
    extra_instruction = """
    Create:
    - Short explanations
    - Bullet points
    - Important highlights
    - Easy-to-read format
    """

# REPORT
elif mode == "Report":
    extra_instruction = """
    Create:
    - Title
    - Introduction
    - Analysis
    - Findings
    - Conclusion
    - Professional formatting
    """

# ---------------- PROMPT ----------------
prompt = f"""
{task}

Topic: {topic}

Resume Information:
{resume_data}

Tone: {tone}

Length: approximately {word_limit} words

{extra_instruction}

Use proper formatting.
Make the content professional and well-structured.
Avoid unnecessary repetition.
"""

# ---------------- GENERATE ----------------
if generate_btn:

    # Empty Topic Check
    if not topic:
        st.warning("Please enter a topic")
        st.stop()

    st.info("Generating LIVE content... ⚡")

    # ---------------- OLLAMA STREAMING ----------------
    stream = ollama.chat(
        model="phi3:mini",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        stream=True
    )

    full_text = ""

    placeholder = st.empty()

    # Live Streaming Output
    for chunk in stream:

        if "message" in chunk:

            content = chunk["message"]["content"]

            full_text += content

            placeholder.write(full_text)

    # ---------------- WORD COUNT ----------------
    word_count = len(full_text.split())

    st.success(f"✅ Document Generated Successfully ({word_count} words)")

    # ---------------- DOCX GENERATION ----------------
    doc = Document()

    doc.add_heading(f"{mode}: {topic}", level=0)

    for line in full_text.split("\n"):

        line = line.strip()

        if not line:
            continue

        # Heading
        if line.startswith("#") or line.isupper():

            doc.add_heading(
                line.replace("#", ""),
                level=1
            )

        # Bullet Points
        elif line.startswith("-") or line.startswith("•"):

            doc.add_paragraph(
                line,
                style="List Bullet"
            )

        # Normal Paragraph
        else:
            doc.add_paragraph(line)

    # Save DOCX
    docx_file = topic.replace(" ", "_") + ".docx"

    doc.save(docx_file)

    # DOCX Download
    with open(docx_file, "rb") as f:

        st.download_button(
            "📥 Download DOCX",
            f,
            file_name=docx_file
        )

    # ---------------- PDF GENERATION ----------------
    pdf_file = topic.replace(" ", "_") + ".pdf"

    pdf = SimpleDocTemplate(pdf_file)

    styles = getSampleStyleSheet()

    style = styles["BodyText"]

    elements = []

    # PDF Title
    elements.append(
        Paragraph(
            f"<b>{mode}: {topic}</b>",
            styles["Title"]
        )
    )

    elements.append(Spacer(1, 12))

    # PDF Content
    for line in full_text.split("\n"):

        line = line.strip()

        if line:

            elements.append(
                Paragraph(line, style)
            )

            elements.append(
                Spacer(1, 6)
            )

    # Build PDF
    pdf.build(elements)

    # PDF Download
    with open(pdf_file, "rb") as f:

        st.download_button(
            "📄 Download PDF",
            f,
            file_name=pdf_file
        )
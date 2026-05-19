import ollama
from docx import Document

# 1. User input
topic = input("Enter topic: ")
pages = input("How many pages (approx): ")

print("Generating content...")

# 2. AI prompt
prompt = f"""
Write a detailed report on {topic}.

Make it long enough for approximately {pages} pages.

Include:
- Introduction
- Key Concepts
- Advantages
- Disadvantages
- Applications
- Future Scope
- Conclusion

Write in structured paragraphs.
"""

# 3. Call Ollama
response = ollama.chat(
    model="phi3:mini",
    messages=[{"role": "user", "content": prompt}]
)

content = response["message"]["content"]

# 4. Create Word file
doc = Document()
doc.add_heading(topic, level=1)
doc.add_paragraph(content)

# 5. Save file
filename = topic.replace(" ", "_") + ".docx"
doc.save(filename)

print("DOCX created successfully:", filename)